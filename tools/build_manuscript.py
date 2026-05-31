"""
manuscript.docx 템플릿(KSME 양식)의 스타일/2단 구조를 보존한 채
음식물 쓰레기 자율 수거 로봇 시제품 논문을 생성한다.

본문의 정량 수치는 tools/experiments_results.json (tools/run_experiments.py 산출)에서
읽어와 채운다. 데이터를 코드에 임의로 박지 않고 실험 결과와 동기화한다.

사용:
    source .venv/bin/activate
    python tools/run_experiments.py      # (먼저) 실험 수치 생성
    python tools/build_manuscript.py     # 논문 생성
출력:
    논문_food_waste_robot.docx
"""
import json
import os
import re
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = "/Users/kyungsbook/Desktop/autoproject"
TEMPLATE = f"{ROOT}/manuscript.docx"
OUTPUT = f"{ROOT}/논문_food_waste_robot.docx"
RESULTS = f"{ROOT}/tools/experiments_results.json"

with open(RESULTS) as f:
    R = json.load(f)
E1, E2, E3, E4 = (R["E1_mission_kpi"], R["E2_inflation"],
                  R["E3_tsp_quality"], R["E4_scalability"])

d = docx.Document(TEMPLATE)

# --- 구역나누기(1단->2단) anchor 찾기 + 본문 비우기 ---
anchor = None
for p in d.paragraphs:
    ppr = p._p.find(qn("w:pPr"))
    if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
        anchor = p
        break
assert anchor is not None
for p in list(d.paragraphs):
    if p._p is not anchor._p:
        p._p.getparent().remove(p._p)
for t in list(d.tables):
    t._tbl.getparent().remove(t._tbl)
for r in list(anchor.runs):
    r._r.getparent().remove(r._r)


# ---------------- helpers ----------------
def tnr(run, bold=None, size=None, italic=None):
    run.font.name = "Times New Roman"
    if bold is not None:
        run.font.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if italic is not None:
        run.font.italic = italic
    return run


CITE = re.compile(r"\{(\d+(?:[,~]\d+)*)\}")


def add_rich(p, text, bold=False):
    pos = 0
    for m in CITE.finditer(text):
        if m.start() > pos:
            tnr(p.add_run(text[pos:m.start()]), bold=bold)
        sup = tnr(p.add_run(f"({m.group(1)})"))
        sup.font.superscript = True
        pos = m.end()
    if pos < len(text):
        tnr(p.add_run(text[pos:]), bold=bold)
    return p


def before(text, style, bold=False, size=None):
    p = anchor.insert_paragraph_before("", style)
    if text:
        tnr(p.add_run(text), bold=bold, size=size)
    return p


def body(text, style="본문 단락:논문용", bold=False):
    p = d.add_paragraph(style=style)
    if text:
        add_rich(p, text, bold=bold)
    return p


def heading(text, style):
    p = d.add_paragraph(style=style)
    tnr(p.add_run(text))
    return p


def set_borders(table):
    tblPr = table._tbl.tblPr
    b = tblPr.find(qn("w:tblBorders"))
    if b is None:
        b = tblPr.makeelement(qn("w:tblBorders"), {})
        tblPr.append(b)
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.makeelement(qn("w:" + e), {})
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        b.append(el)


def add_table(caption, header, rows):
    cap = d.add_paragraph(style="Normal")
    tnr(cap.add_run(caption), size=10)
    t = d.add_table(rows=1 + len(rows), cols=len(header))
    set_borders(t)
    for j, htxt in enumerate(header):
        tnr(t.rows[0].cells[j].paragraphs[0].add_run(htxt), size=9, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            tnr(t.rows[i].cells[j].paragraphs[0].add_run(str(val)), size=9)
    return t


def figure_ph(caption):
    ph = d.add_paragraph(style="그림 번호/설명")
    tnr(ph.add_run("[ Figure placeholder — 이미지 삽입 필요 ]"), italic=True)
    cap = d.add_paragraph(style="그림 번호/설명")
    tnr(cap.add_run(caption))


# ============================================================
# 섹션 0 : 제목 / 저자 / 키워드 / 초록  (1단)
# ============================================================
before("공동주택 음식물 쓰레기 자율 수거를 위한 저비용 임베디드 시제품 개발",
       "한글 논문제목", bold=True)
before("", "Normal")
before("이경석 · 이상윤 · 전민제", "저자 이름")
before("한양대학교 기계공학부", "저자 소속")
before("", "Normal")
before("Development of a Low-Cost Embedded Prototype for Autonomous "
       "Food Waste Collection in Residential Complexes",
       "영어 논문 제목", bold=True)
before("", "Normal")
before("Gyeongseok Lee, Sangyun Lee and Minje Jeon", "영어 저자 이름")
before("School of Mechanical Engineering, Hanyang University.", "저자 소속")
before("", "Normal")
before("Key Words: Autonomous Mobile Robot(자율주행 로봇), Food Waste Collection"
       "(음식물 쓰레기 수거), Embedded System(임베디드 시스템), A* Path Planning"
       "(A* 경로 계획), QR Code Localization(QR 코드 측위)", "Normal", size=10)

kor_abs = (
    "초록: 본 논문은 공동주택 단지에서 3 L 규모의 음식물 쓰레기통을 자율적으로 수거하는 "
    "저비용 임베디드 로봇 시제품과, 이를 일관되게 검증하기 위한 3계층 개발 환경을 제안한다. "
    "제안 플랫폼은 웹 기반 2D 시뮬레이션, Webots 3D 물리 시뮬레이션, 실물 로봇이 동일한 미션 "
    "로직을 공유하며 실시간으로 동기화되는 구조를 갖는다. 시제품은 Raspberry Pi 4와 Arduino Mega "
    "기반의 이원화 제어 구조에, 고가의 LiDAR·GPS 없이 초음파·관성·영상 센서만을 결합하였고, A* 경로 탐색, 최근접 이웃 "
    "방문 순서 최적화, QR 코드 측위, 유한 상태머신 기반 수거 시퀀스로 빈 탐색·접근·수거·배출을 "
    "수행한다. 마이크로컨트롤러 수준의 독립 안전 정지 로직으로 충돌과 통신 두절에 대응한다. "
    "40×30 격자 시제품 환경의 시뮬레이션 실험을 통해 수거 미션 시간, 장애물 팽창 비용의 효과, "
    "방문 순서 최적화 품질, 경로 탐색 확장성을 정량적으로 분석하였다."
)
before(kor_abs, "바탕글")

eng_abs = (
    "Abstract: This paper presents a low-cost embedded prototype robot for "
    "autonomously collecting 3 L food-waste bins in residential apartment "
    "complexes, together with a three-layer development environment for "
    "consistent verification. The platform lets a web-based 2D simulation, a "
    "Webots 3D physics simulation, and the physical robot share an identical "
    "mission logic synchronized in real time. The prototype combines a "
    "Raspberry Pi 4 and an Arduino Mega in a two-tier control architecture "
    "with ultrasonic, inertial, and vision sensors but without LiDAR or GPS, "
    "and performs bin search, "
    "approach, pickup, and disposal through A* path planning, nearest-neighbor "
    "visit-order optimization, QR-code localization, and a finite-state "
    "mission sequence. An independent microcontroller-level safety-stop logic "
    "handles collision and communication loss. Using a 40x30 grid prototype "
    "environment, we quantitatively analyze mission time, the effect of "
    "obstacle inflation cost, the quality of visit-order optimization, and the "
    "scalability of the path planner through simulation experiments."
)
before(eng_abs, "abstract(내용)")

# ============================================================
# 섹션 1 : 본문  (2단)
# ============================================================
heading("- 기호 설명 -", "기호설명")
body("g(n) : 출발 노드로부터 노드 n 까지의 누적 이동 비용")
body("h(n) : 노드 n 으로부터 목표까지의 휴리스틱 추정 비용")
body("f(n) : A* 의 평가 함수, f(n) = g(n) + h(n)")
body("r_infl : 장애물 주변 비용 팽창 반경 [cells]")
body("v : 로봇 선속도 [cells/s]")

# --- 1. 서론 ---
heading("1. 서 론", "논문의 각 장")
heading("1.1 연구 배경", "논문의 각 절")
body("공동주택 단지에서 음식물 쓰레기의 수거는 위생, 악취, 인력 비용 측면에서 지속적인 부담으로 "
     "지적되어 왔다. 대부분의 단지는 고정된 집하 공간에 주민이 직접 배출하는 방식을 따르며, 수거 "
     "동선과 시점이 사람의 노동에 의존한다. 이러한 반복적·정형적 작업은 자율주행 로봇을 통한 자동화의 "
     "대표적 적용 대상이다.")
body("그러나 상용 자율주행 로봇은 고가의 LiDAR와 산업용 제어기를 전제로 하여, 소규모 단지나 시범 "
     "도입 단계에서 비용 장벽이 크다. 또한 알고리즘 검증과 실물 하드웨어 동작 사이의 간극이 커, 개발 "
     "과정에서 반복적인 현장 시험이 요구되며 이는 개발 비용과 위험을 증가시킨다.")
heading("1.2 관련 연구", "논문의 각 절")
body("격자 지도 기반의 자율주행 경로 계획에서는 A* 알고리즘{1}이 최적성과 효율의 균형으로 널리 "
     "활용되며, ROS 2의 Nav2 프레임워크{2}는 이를 전역 계획기와 비용 지도(costmap) 기반의 지역 "
     "계획기로 통합한다. 저비용 측위를 위해서는 QR 코드와 같은 인공 표식 기반 기법{3}이 제안되어 "
     "왔고, 보행자·장애물 인지에는 YOLO{4} 계열의 실시간 객체 검출이 사용된다. 본 연구는 이러한 "
     "구성요소를 저비용 임베디드 하드웨어에 통합하고, 시뮬레이션과 실물을 일관되게 검증하는 환경에 "
     "초점을 둔다.")
heading("1.3 연구 목표 및 기여", "논문의 각 절")
body("본 연구의 기여는 다음과 같다. 첫째, 저비용 임베디드 부품만으로 음식물 쓰레기 수거 로봇 시제품을 "
     "구현하였다. 둘째, 웹 2D 시뮬레이션·Webots 3D 시뮬레이션·실물 로봇이 동일 미션 로직을 공유하는 "
     "3계층 검증 환경을 제안하였다. 셋째, ROS 2 Nav2 이식을 염두에 둔 A* 경로 탐색·안전 구조를 "
     "설계하고, 시제품 환경에서 미션 성능과 알고리즘 특성을 정량적으로 분석하였다.")

# --- 2. 시스템 구성 ---
heading("2. 시스템 구성", "논문의 각 장")
heading("2.1 3계층 검증 아키텍처", "논문의 각 절")
body("제안 플랫폼은 Fig. 1과 같이 세 계층으로 구성된다. 웹 2D 시뮬레이션 계층은 알고리즘의 빠른 "
     "반복 검증을, Webots 3D 시뮬레이션 계층은 센서·모터 동역학을 포함한 물리 검증을, 실물 로봇 "
     "계층은 실제 하드웨어 동작을 담당한다. 세 계층은 WebSocket을 통해 5 Hz로 상태를 동기화하여 "
     "동일한 미션 로직이 일관되게 동작하도록 하며, 한 계층에서 검증된 로직을 다른 계층으로 위험 없이 "
     "이전할 수 있게 한다. 다만 측위 수단은 계층별로 다른데, Webots 3D 시뮬레이션 계층은 GPS·Compass를 "
     "지면 실측(ground-truth) 값으로 제공받아 알고리즘 자체를 검증하는 반면, 실물 로봇 계층은 동일한 미션 "
     "로직을 GPS 없이 QR 코드 측위와 관성 센서 기반 추측 항법(dead reckoning)으로 구동한다.")
figure_ph("Fig. 1 Three-layer verification architecture synchronizing the web "
          "2D simulation, the Webots 3D simulation, and the physical robot")
add_table("Table 1 Three-layer verification environment",
          ["Layer", "Purpose", "Verifies"],
          [["Web 2D simulation", "Fast algorithm iteration", "Path planning, mission logic"],
           ["Webots 3D simulation", "Physical validation", "Sensor / motor dynamics"],
           ["Physical robot", "Real-world operation", "Hardware behavior, safety"]])
heading("2.2 하드웨어 설계", "논문의 각 절")
body("시제품의 주요 부품 구성은 Table 2와 같다. 제어부는 비전·미션 처리를 담당하는 Raspberry Pi 4와 "
     "실시간 센서·구동 제어를 담당하는 Arduino Mega로 이원화하였다. 고수준 인지·계획은 Linux 기반 "
     "Raspberry Pi가, 결정론적 실시간 제어와 안전 정지는 마이크로컨트롤러가 담당함으로써 안전성과 "
     "유연성을 동시에 확보한다. 초기 설계에서는 옥외 자율주행을 위한 GPS 측위 모듈을 포함할 계획이었으나, "
     "시제품 단계의 예산 제약과 단지 내 고층 건물에 의한 위성 신호 음영을 고려하여 GPS를 배제하고, 초음파 "
     "센서·Raspberry Pi 카메라·USB 웹캠·관성 센서를 결합한 저비용 센서 구성으로 대체하였다. 이는 고가의 "
     "측위 장비 없이도 단지 환경에서 자율 수거가 가능함을 보이려는 본 연구의 지향과 일치한다. 두 제어기는 "
     "USB 시리얼로 연결되며, 로직용 7.4 V 배터리와 구동용 "
     "12 V 배터리를 분리한 듀얼 배터리 전원 구조(Fig. 2)를 채택해 모터 노이즈의 제어부 유입을 "
     "억제하였다.")
add_table("Table 2 Main hardware components of the prototype",
          ["Category", "Component"],
          [["Controller", "Raspberry Pi 4 (4 GB) + Arduino Mega 2560"],
           ["Vision", "RPi Camera Module 3 + USB webcam"],
           ["Sensing", "MPU-9250 IMU + HC-SR04 ultrasonic x5"],
           ["Actuation", "L298N x2 + DC geared motor x2 + MG996R servo"],
           ["Power", "2S LiPo 7.4 V (logic) + 12 V LiPo (drive)"]])
figure_ph("Fig. 2 Dual-battery power architecture separating the 7.4 V logic "
          "supply from the 12 V drive supply")
heading("2.3 소프트웨어 구성", "논문의 각 절")
body("백엔드는 Python FastAPI와 비동기 SQLAlchemy로 구현되어 미션·로봇·빈 상태를 관리하고 "
     "WebSocket으로 시뮬레이션 상태를 브로드캐스트한다. 실물 로봇에서는 Raspberry Pi가 미션 "
     "상태머신과 비전 처리를, Arduino가 모터·센서 제어를 담당하며, 두 제어기는 115200 bps의 JSON "
     "라인 프로토콜로 통신한다. 동일한 미션 로직 모듈을 세 계층이 공유하도록 설계하여 코드 중복과 "
     "검증 불일치를 최소화하였다.")

# --- 3. 자율 수거 알고리즘 ---
heading("3. 자율 수거 알고리즘", "논문의 각 장")
heading("3.1 경로 탐색", "논문의 각 절")
body("전역 경로 탐색에는 격자 지도 상의 A* 알고리즘{1}을 사용한다. 평가 함수는 f(n) = g(n) + h(n) "
     "이며, 8방향 이동을 허용하고 대각 이동 비용을 √2 로 부여한다. 휴리스틱 h(n) 으로는 8방향 이동에 "
     "일관적인 옥타일 거리(octile distance)를 사용한다. 장애물 주변에는 팽창 반경 r_infl 을 적용해 "
     "비용을 가중함으로써 ROS 2 Nav2의 InflationLayer 동작을 모사하였으며, 이는 향후 Nav2 NavFn "
     "계획기로의 이식을 용이하게 하기 위한 의도적 설계이다.")
heading("3.2 방문 순서 최적화", "논문의 각 절")
body("다수의 빈을 수거할 때의 방문 순서는 최근접 이웃(nearest-neighbor) 휴리스틱으로 결정한다. 이는 "
     "외판원 문제(TSP)의 근사 해법으로, 계산 비용이 낮아 임베디드 환경에 적합하다. 현재 위치에서 가장 "
     "가까운 미방문 빈을 반복적으로 선택하며, 마지막에 집하장으로 복귀하는 경로를 구성한다.")
heading("3.3 비전 인식", "논문의 각 절")
body("각 음식물 쓰레기통에는 고유 QR 코드{3}를 부착하여 빈 식별과 접근 정렬에 활용한다. 로봇은 QR "
     "경계 상자의 화면 내 위치를 카메라 수평 화각에 매핑하여 빈에 대한 방위각을 추정하고, 전방 초음파 "
     "거리로 접근 거리를 판단한다. 후방 카메라의 객체 검출{4}과 측·후방 초음파를 융합하여 보행자 및 "
     "장애물을 감지한다.")
heading("3.4 미션 상태머신 및 안전 시스템", "논문의 각 절")
body("수거 미션은 Fig. 3과 같이 IDLE → NAV_TO_BIN → APPROACH → PICKUP → NAV_TO_DEPOT → DROP "
     "의 유한 상태머신으로 구성된다. APPROACH 단계에서 목표 QR이 일치하고 30 cm 이내로 접근하면 "
     "롤러를 구동해 수거(PICKUP)하고, 집하장으로 이동한 뒤 역방향 롤러로 배출(DROP)한다. 한편 "
     "Arduino는 상위 제어기의 명령보다 우선하는 독립 안전 정지 로직을 수행하여, 전방 장애물이 15 cm "
     "미만이고 전진 중이면 즉시 정지하고, 500 ms 이상 명령이 없으면 워치독에 의해 정지한다. 안전 정지 "
     "발생 시 상위 제어기는 후진·회전으로 회피한 뒤 미션을 재개한다.")
figure_ph("Fig. 3 Finite-state machine of the collection mission")

# --- 4. 실험 및 분석 ---
heading("4. 실험 및 분석", "논문의 각 장")
heading("4.1 실험 환경", "논문의 각 절")
body(f"실험은 시제품 테스트 환경을 모사한 {R['meta']['grid']} 격자(장애물 셀 "
     f"{R['meta']['n_obstacles']}개, 건물 4동·놀이터·주차장 포함)에서 수행하였다. 빈 4개와 로봇 "
     "2대, 중앙 집하장을 배치하였으며, 경로 탐색은 기본 팽창 반경 r_infl = 2, 시제품 공칭 속도 "
     f"v = {R['meta']['robot_speed_cells_s']} cells/s, 빈당 수거 시간 "
     f"{int(R['meta']['pickup_time_s'])} s 를 적용하였다. 모든 수치는 실제 백엔드 알고리즘 구현을 "
     "그대로 실행하여 산출하였으며 난수 시드를 고정해 재현성을 확보하였다. 본 절의 결과는 "
     "시뮬레이션 기반 분석으로, 실물 하드웨어 계측은 향후 과제로 남긴다.")

heading("4.2 수거 미션 성능", "논문의 각 절")
s4 = E1["single_robot_all4"]
dr = E1["dual_robot"]
body(f"단일 로봇이 4개 빈을 모두 수거하는 경우, 최근접 이웃 방문 순서에 따른 총 이동 거리는 "
     f"{s4['total_distance']} cells, 추정 미션 시간은 {s4['est_time_sec']} s 였다. 로봇 2대가 "
     "좌·우 구역을 분담하는 경우 두 로봇의 작업을 병렬화할 수 있으며, 더 늦게 끝나는 로봇 기준의 "
     f"makespan은 {dr['makespan_sec']} s 로, 단일 로봇 대비 약 "
     f"{dr['throughput_improvement_pct']}% 의 시간 단축을 보였다(Table 3). 두 로봇 모두 작업 후 "
     "중앙 집하장으로 복귀해야 하므로 복귀 구간이 전체 시간에서 큰 비중을 차지하며, 이는 다중 로봇 "
     "병렬화 이득을 제한하는 요인으로 분석된다.")
add_table("Table 3 Collection mission performance (simulation)",
          ["Configuration", "Distance [cells]", "Time [s]"],
          [["Single robot (4 bins)", s4["total_distance"], s4["est_time_sec"]],
           ["Dual robot - A (left)", dr["robotA_left"]["total_distance"], dr["robotA_left"]["est_time_sec"]],
           ["Dual robot - B (right)", dr["robotB_right"]["total_distance"], dr["robotB_right"]["est_time_sec"]],
           ["Dual robot (makespan)", "-", dr["makespan_sec"]]])

heading("4.3 장애물 팽창 비용의 효과", "논문의 각 절")
e2lo, e2hi = E2[0], E2[-1]
body(f"팽창 반경 r_infl 을 0 에서 4 까지 변화시키며 단일 로봇 미션 경로를 분석하였다(Table 4). "
     f"r_infl 이 증가할수록 경로의 평균 장애물 이격 거리는 {e2lo['mean_clearance_cells']} 에서 "
     f"{e2hi['mean_clearance_cells']} cells 로 증가하여 더 안전한 여유 공간을 확보하는 반면, 총 "
     f"경로 길이는 {e2lo['total_path_len']} 에서 {e2hi['total_path_len']} cells 로 늘고 A* "
     f"질의당 계산 시간도 {e2lo['avg_astar_ms']} 에서 {e2hi['avg_astar_ms']} ms 로 증가한다. "
     "최소 이격 거리가 1 cell 로 일정한 것은 빈과 집하장이 건물·구조물에 인접해 있어 출발·도착 "
     "구간에서 근접이 불가피하기 때문이며, 이 trade-off를 고려해 본 연구는 r_infl = 2 를 기본값으로 "
     "채택하였다.")
add_table("Table 4 Effect of inflation radius on the single-robot mission path",
          ["r_infl", "Path len [cells]", "Mean clear. [cells]", "A* time [ms]"],
          [[r["inflation_radius"], r["total_path_len"], r["mean_clearance_cells"], r["avg_astar_ms"]]
           for r in E2])

heading("4.4 방문 순서 최적화 품질", "논문의 각 절")
real = next(x for x in E3 if x["case"] == "real_4bin")
rand = [x for x in E3 if x["case"] == "random"]
g8 = next(x for x in rand if x["N"] == 8)
body(f"실제 4-bin 레이아웃에서 최근접 이웃 휴리스틱은 완전 탐색 최적해와 동일한 경로를 산출하여 "
     f"최적성 격차가 {real['gap_pct']}% 였다. 빈 개수를 4 에서 8 까지 늘린 무작위 인스턴스 "
     f"({rand[0]['instances']}회 평균)에서는 평균 격차가 N=4 의 {rand[0]['mean_gap_pct']}% 에서 "
     f"N=8 의 {g8['mean_gap_pct']}% 로 증가하였다(Table 5). 이는 소규모 빈 배치에서는 최근접 이웃 "
     "휴리스틱이 최적에 근접하나, 빈 수가 많아지는 대규모 단지에서는 2-opt 등 개선된 TSP 기법이 "
     "필요함을 시사한다.")
add_table("Table 5 Optimality gap of the nearest-neighbor visit order",
          ["Case", "N", "Mean gap [%]", "Max gap [%]"],
          [["Real layout", real["N"], real["gap_pct"], real["gap_pct"]]] +
          [["Random", r["N"], r["mean_gap_pct"], r["max_gap_pct"]] for r in rand])

heading("4.5 경로 탐색 확장성", "논문의 각 절")
e4lo, e4hi = E4[0], E4[-1]
body(f"풀스케일 단지로의 확장 가능성을 확인하기 위해 빈 격자 크기를 {e4lo['grid']} 에서 "
     f"{e4hi['grid']} 까지 늘리며 대각선 경로 질의의 계산 시간을 측정하였다(Table 6). 격자 셀 수가 "
     f"{e4lo['cells']} 에서 {e4hi['cells']} 로 약 23배 증가하는 동안 질의 시간은 "
     f"{e4lo['avg_time_ms']} 에서 {e4hi['avg_time_ms']} ms 로 증가하여, 가장 큰 "
     f"{e4hi['grid']} 격자에서도 50 ms 미만의 실시간 성능을 보였다. 이는 제안 경로 탐색이 실제 "
     "아파트 단지 규모에서도 실시간으로 동작 가능함을 의미한다.")
add_table("Table 6 Path-planning scalability over grid size",
          ["Grid", "Cells", "Path [cells]", "Time [ms]"],
          [[r["grid"], r["cells"], r["path_cells"], r["avg_time_ms"]] for r in E4])

# --- 5. 결론 ---
heading("5. 결 론", "논문의 각 장")
body("본 논문은 공동주택 음식물 쓰레기 자율 수거를 위한 저비용 임베디드 시제품과, 웹·Webots·실물의 "
     "3계층 검증 환경을 제안하였다. Raspberry Pi와 Arduino 기반의 이원화 제어, A* 경로 탐색과 "
     "최근접 이웃 방문 순서 최적화, QR 측위, 마이크로컨트롤러 수준의 독립 안전 로직을 통합하였으며, "
     "시제품 환경의 시뮬레이션 실험으로 미션 성능과 알고리즘 특성을 정량적으로 분석하였다.")
body("실험을 통해 (1) 2대의 로봇 분담이 단일 로봇 대비 미션 시간을 단축하나 집하장 복귀 오버헤드가 "
     "그 이득을 제한하고, (2) 장애물 팽창 비용이 경로 안전 여유와 계산 비용 사이의 trade-off를 "
     "형성하며, (3) 최근접 이웃 휴리스틱은 소규모 배치에서 최적에 근접하나 빈 수 증가에 따라 격차가 "
     "커지고, (4) 경로 탐색이 단지 규모 격자에서도 실시간으로 동작함을 확인하였다.")
body("본 연구의 한계로, 제시한 정량 결과는 시뮬레이션 기반이며 실물 하드웨어 계측은 포함하지 않는다. "
     "또한 현재 시제품의 위치 추정은 관성 센서와 시간 적분 기반의 추측 항법에 의존하여 누적 오차가 "
     "존재하고, 비전 인식은 사전 학습된 검출기에 의존한다. 향후 (1) 실물 로봇의 수거 성공률·위치 "
     "정확도 계측, (2) ROS 2 Nav2 및 EKF 기반 측위로의 이식, (3) 음식물 쓰레기통 전용 객체 검출기의 "
     "데이터셋 구축 및 학습, (4) 다중 로봇 작업 분배와 집하장 배치 최적화를 통해 이를 개선할 계획이다.")

# --- 후기 ---
heading("후 기", "후기(제목)")
body("본 연구는 한양대학교 기계공학부 종합설계 교과의 일환으로 수행되었다.")

# --- 참고문헌 ---
heading("참고문헌", "참고문헌(제목)")
refs = [
    '(1) Hart, P. E., Nilsson, N. J. and Raphael, B., 1968, "A Formal Basis for '
    'the Heuristic Determination of Minimum Cost Paths," IEEE Transactions on '
    'Systems Science and Cybernetics, Vol. 4, No. 2, pp. 100~107.',
    '(2) Macenski, S., Martín, F., White, R. and Clavero, J. G., 2020, "The '
    'Marathon 2: A Navigation System," Proc. IEEE/RSJ Int. Conf. on Intelligent '
    'Robots and Systems, pp. 2718~2725.',
    '(3) Garrido-Jurado, S., Muñoz-Salinas, R., Madrid-Cuevas, F. J. and '
    'Marín-Jiménez, M. J., 2014, "Automatic Generation and Detection of Highly '
    'Reliable Fiducial Markers under Occlusion," Pattern Recognition, Vol. 47, '
    'No. 6, pp. 2280~2292.',
    '(4) Redmon, J., Divvala, S., Girshick, R. and Farhadi, A., 2016, "You Only '
    'Look Once: Unified, Real-Time Object Detection," Proc. IEEE Conf. on '
    'Computer Vision and Pattern Recognition, pp. 779~788.',
]
for r in refs:
    p = d.add_paragraph(style="참고문헌(내용)")
    tnr(p.add_run(r), size=9)

d.save(OUTPUT)
print("saved:", OUTPUT)
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables), "sections:", len(d.sections))
